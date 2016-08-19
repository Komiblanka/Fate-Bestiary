import urllib2
import re
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
import os, sys
from urllib2 import HTTPError
from HTMLParser import HTMLParser

class creature:
	Name = ""
	High_Concept = ""
	Trouble = ""
	Other_Aspects = ""
	Scale = ""
	Skills = []
	Stunts = {}
	Extras = {}
	Physical_Stress = ""
	Mental_Stress = ""
	Size = ""
	Optional = ""

	text = ""
	image = ""
	imagelink = ""
	imagename = ""

	def __init__(self):
		self.Stunts = {}
		self.Extras = {}


	def printcreature(self):
		print "Name: " + self.Name
		print "High Concept: " + self.High_Concept
		print "Trouble: " + self.Trouble
		print "Other Aspects: " + self.Other_Aspects
		print "Scale: " + self.Scale
		print "Skills: "
		for skill in self.Skills:
			print "\t" + skill
		print "Stunts: "
		for name, stunt in self.Stunts.iteritems():
			print "\t" + name + ": " + stunt
		print "Extras: "
		for extra in self.Extras:
			print "\t" + extra
		print "Physical Stress: " + self.Physical_Stress
		print "Mental Stress: " + self.Mental_Stress
		print "Size: " + self.Size
		print "Optional: " + self.Optional
		#print "*************************************"
		#print self.text
	def save_image(self):
		localFile = open(self.imagename, 'wb')
		localFile.write(self.image.read())
		localFile.close()

	def set_attributes(self, link):

		response = urllib2.urlopen(link)
		html = response.read()
		soup = BeautifulSoup(html, "html.parser")

		parser = HTMLParser()



		self.Name = soup.find("a", {"name":"creaturename"}).string
		self.text = soup.table.getText().replace("High Concept", "\nHigh Concept")

		print "Getting creature: " + self.Name;

		link = soup.table.img["src"]

		self.imagename = link.split("/")[-1]
		try:
			self.image = urllib2.urlopen(link)
		except HTTPError:
			print "No image for " + self.Name
			self.imagename = "none"

		p = re.compile('High Concept:(.*)')
		self.High_Concept = re.findall(p, self.text)[0]

		p = re.compile('Trouble:(.*)')
		self.Trouble = re.findall(p, self.text)[0]

		p = re.compile('Other Aspects:(.*)')
		self.Other_Aspects = re.findall(p, self.text)[0]

		p = re.compile('Scale:(.*)')
		self.Scale = re.findall(p, self.text)[0]

		p = re.compile('Skills:(.*)')
		self.Skills = re.findall(p, str(soup.table))[0].split("top\">")[1].split("</")[0].split("<br/>")

		p = re.compile('Stunts:(.*)')
		stunt_text = re.findall(p, str(parser.unescape(soup.table)))[0].split("top\">")[1].split("</td")[0].split("<br/>")

		for stunt in stunt_text:
			if "</strong>" in stunt:
				self.Stunts.update({stunt.split("</strong>")[0].split("<strong>")[1] : stunt.split("</strong>")[1].replace("\xe2\x80\x94","--")})
			else:
				self.Stunts.update({"" : stunt.replace("\xe2\x80\x94","--")})


		p = re.compile('Extras:(.*)')
		extras_text = re.findall(p, str(soup.table))[0].split("top\">")[1].split("</td")[0].split("<br/>")

		for extra in extras_text:
			if "</strong>" in extra:
				self.Extras.update({extra.split("</strong>")[0].split("<strong>")[1] : extra.split("</strong>")[1].replace("\xe2\x80\x94","--")})

		p = re.compile('Physical Stress:(.*)')
		self.Physical_Stress = re.findall(p, self.text)[0]

		p = re.compile('Mental Stress:(.*)')
		self.Mental_Stress = re.findall(p, self.text)[0]

		p = re.compile('Size:(.*)')
		self.Size = re.findall(p, self.text)[0]

		p = re.compile('Optional:(.*)')
		texto = re.findall(p, self.text)
		if len(texto) > 0:
			self.Optional = texto[0].replace(".", ". ")
def get_creatures(links):
	monsters = []
	for link in links:
		monster = creature()
		monster.set_attributes(link)
		monsters.append(monster)
	return monsters

def get_links():
	response = urllib2.urlopen('http://inkwellideas.com/fate_creatures/')
	html = response.read()
	soup = BeautifulSoup(html, "html.parser")
	links = []

	for link in soup.find_all('td'):
		links.append("http://inkwellideas.com/fate_creatures/" + link.a.get('href'))
	return links
def write_monster(monsters):

	document = Document()

	document.add_heading('Fate Monsters', 0)

	for monster in monsters:
		print "Writing creature: " + monster.Name
		document.add_heading(monster.Name, 1)

		if monster.imagename != "none":
			monster.save_image()
			document.add_picture(monster.imagename, width=Inches(2))

		p = document.add_paragraph()
		p.add_run("High Concept: ").bold=True
		p.add_run(monster.High_Concept)
		p = document.add_paragraph()
		p.add_run("Trouble: ").bold=True
		p.add_run(monster.Trouble)
		p = document.add_paragraph()
		p.add_run("Other Aspects: ").bold=True
		p.add_run(monster.Other_Aspects)
		p = document.add_paragraph()
		p.add_run("Scale: ").bold=True
		p.add_run(monster.Scale)
		p = document.add_paragraph()

		p.add_run("Skills:").bold=True
		for skill in monster.Skills:
			p = document.add_paragraph(skill, style='ListBullet')
			p.paragraph_format.left_indent = Inches(0.5)
		parser = HTMLParser()

		p = document.add_paragraph()
		p.add_run("Stunts: ").bold=True
		for name, description in monster.Stunts.iteritems():
			p = document.add_paragraph(style='ListBullet'); p.paragraph_format.left_indent = Inches(0.5)
			p.add_run(name + ": ").bold=True; p.add_run(description);



		p = document.add_paragraph()
		p.add_run("Extras: ").bold=True
		for name, description in monster.Extras.iteritems():
			p = document.add_paragraph(style='ListBullet'); p.paragraph_format.left_indent = Inches(0.5)
			p.add_run(name + ": ").bold=True; p.add_run(description);


		p = document.add_paragraph()
		p.add_run("Physical Stress: ").bold=True
		p.add_run(monster.Physical_Stress)
		p = document.add_paragraph()
		p.add_run("Mental Stress: ").bold=True
		p.add_run(monster.Mental_Stress)
		p = document.add_paragraph()
		p.add_run("Size: ").bold=True
		p.add_run(monster.Size)

		p = document.add_paragraph()
		p.add_run("Optional: ").bold=True
		p.add_run(monster.Optional)

		document.add_page_break()

		if monster.imagename != "none":
			os.remove(monster.imagename)


	document.save('monstertest.docx')

print "Getting links\n"
links = get_links()
print "Getting monsters\n"
monsters = get_creatures(links)
print "Writting document\n"
write_monster(monsters)


#monster = creature()
#monster.set_attributes("http://inkwellideas.com/fate_creatures/bear_grizzly.shtml")
#monsters = []
#monsters.append(monster)
##print monster.Optional
#write_monster(monsters)